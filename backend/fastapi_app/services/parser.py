import io
import PyPDF2
import docx
from typing import Union
import logging

logger = logging.getLogger(__name__)


class DocumentParser:
    """Service for extracting text from various document formats"""
    
    def __init__(self):
        self.supported_types = {
            'application/pdf': self._extract_from_pdf,
            'application/msword': self._extract_from_doc,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_from_docx,
            'text/plain': self._extract_from_txt
        }
    
    def extract_text(self, file_content: bytes, content_type: str) -> str:
        """
        Extract text from document based on content type
        """
        if content_type not in self.supported_types:
            raise ValueError(f"Unsupported content type: {content_type}")
        
        try:
            extractor = self.supported_types[content_type]
            text = extractor(file_content)
            return self._clean_text(text)
        except Exception as e:
            logger.error(f"Error extracting text from {content_type}: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF"""
        text = ""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        except Exception as e:
            logger.error(f"Error reading PDF: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
        
        return text
    
    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            doc_file = io.BytesIO(file_content)
            doc = docx.Document(doc_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                    
        except Exception as e:
            logger.error(f"Error reading DOCX: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
        
        return text
    
    def _extract_from_doc(self, file_content: bytes) -> str:
        """Extract text from DOC (legacy Word format)"""
        # For DOC files, we might need to use python-docx2txt or antiword
        # For now, we'll return an error message
        raise ValueError("DOC format not fully supported. Please convert to DOCX or PDF.")
    
    def _extract_from_txt(self, file_content: bytes) -> str:
        """Extract text from plain text file"""
        try:
            # Try UTF-8 first, then fallback to latin-1
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                text = file_content.decode('latin-1', errors='ignore')
            
            return text
        except Exception as e:
            logger.error(f"Error reading text file: {str(e)}")
            raise ValueError(f"Failed to extract text from text file: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:  # Skip empty lines
                cleaned_lines.append(line)
        
        # Join lines with single newlines
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Remove multiple consecutive spaces
        import re
        cleaned_text = re.sub(r' +', ' ', cleaned_text)
        
        return cleaned_text
    
    def get_document_metadata(self, file_content: bytes, content_type: str) -> dict:
        """Extract metadata from document"""
        metadata = {
            'content_type': content_type,
            'size_bytes': len(file_content),
            'page_count': None,
            'word_count': None,
            'char_count': None
        }
        
        try:
            text = self.extract_text(file_content, content_type)
            metadata['word_count'] = len(text.split())
            metadata['char_count'] = len(text)
            
            if content_type == 'application/pdf':
                pdf_file = io.BytesIO(file_content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                metadata['page_count'] = len(pdf_reader.pages)
                
        except Exception as e:
            logger.warning(f"Could not extract metadata: {str(e)}")
        
        return metadata
